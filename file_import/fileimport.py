import os
import logging
import os
import logging
from typing import Dict, Any, Optional
import tempfile
import whisper
from kreuzberg import extract_file_sync, ExtractionConfig, TesseractConfig, PaddleOCRConfig, EasyOCRConfig, PSMMode
import subprocess
from docx import Document
from pptx import Presentation
import pdfplumber
from pathlib import Path
from paddleocr import PPStructureV3, PaddleOCRVL
import datetime
from PIL import Image
import cv2
from openai import OpenAI

model = whisper.load_model("turbo")


class FileParser:
    """
    文件解析与内容提取子模块
    根据文件类型调用相应的解析器提取文本内容和元数据
    """

    def __init__(self, deepseek_api_key: str = None):
        self.logger = logging.getLogger(__name__)
        self.supported_formats = {
            'pdf': self._parse_pdf,
            'ppt': self._parse_ppt,
            'pptx': self._parse_ppt,
            'docx': self._parse_docx,
            'video': self._parse_video,
            'mp4': self._parse_video,
            'avi': self._parse_video,
            'mov': self._parse_video,
            'mkv': self._parse_video,
            'jpg': self._parse_image,
            'jpeg': self._parse_image,
            'png': self._parse_image
        }

        # 初始化DeepSeek客户端
        self.deepseek_client = None
        if deepseek_api_key:
            try:
                self.deepseek_client = OpenAI(
                    api_key=deepseek_api_key,
                    base_url="https://api.deepseek.com/v1"
                )
                self.logger.info("DeepSeek API客户端初始化成功")
            except Exception as e:
                self.logger.error(f"DeepSeek API客户端初始化失败: {str(e)}")

    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        提取文件的基本元数据

        Args:
            file_path: 文件路径

        Returns:
            Dict包含文件的元数据信息
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"文件不存在: {file_path}")

            # 获取文件基本信息
            file_stat = os.stat(file_path)
            file_name = os.path.basename(file_path)
            file_extension = Path(file_path).suffix.lower()
            file_type = file_extension[1:] if file_extension.startswith('.') else file_extension

            # 基本元数据
            metadata = {
                'file_name': file_name,
                'file_type': file_type,
                'file_format': file_extension,
                'file_size': file_stat.st_size,  # 字节
                'file_size_mb': round(file_stat.st_size / (1024 * 1024), 2),  # MB
                'creation_time': datetime.datetime.fromtimestamp(file_stat.st_ctime).strftime('%Y-%m-%d %H:%M:%S'),
                'modification_time': datetime.datetime.fromtimestamp(file_stat.st_mtime).strftime('%Y-%m-%d %H:%M:%S'),
                'access_time': datetime.datetime.fromtimestamp(file_stat.st_atime).strftime('%Y-%m-%d %H:%M:%S')
            }

            return metadata

        except Exception as e:
            self.logger.error(f"提取文件元数据错误: {str(e)}")
            raise

    def extract_content(self, file_path: str, file_type: str, generate_course_metadata: bool = False) -> Dict[str, Any]:
        """
        提取文件内容的入口函数

        Args:
            file_path: 文件路径
            file_type: 文件类型
            generate_course_metadata: 是否生成课程元数据

        Returns:
            Dict包含提取的内容和元数据
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        file_type_lower = file_type.lower()

        if file_type_lower not in self.supported_formats:
            raise ValueError(f"不支持的文件类型: {file_type}")

        try:
            # 提取基本元数据
            metadata = self.extract_metadata(file_path)

            # 根据文件类型提取特定元数据
            specific_metadata = self._extract_specific_metadata(file_path, file_type_lower)
            metadata.update(specific_metadata)

            # 调用相应的解析器提取内容
            content_text = self.supported_formats[file_type_lower](file_path)

            # 生成课程元数据（如果需要）
            course_metadata = {}
            if generate_course_metadata and content_text:
                try:
                    if self.deepseek_client:
                        course_metadata = self._generate_course_metadata(content_text)
                    else:
                        self.logger.warning("DeepSeek客户端未初始化，跳过课程元数据生成")
                except Exception as e:
                    self.logger.warning(f"生成课程元数据失败: {str(e)}")
                    # 即使失败也返回空的课程元数据，确保数据结构一致
                    course_metadata = self._get_empty_course_metadata()

            return {
                'content_text': content_text,
                'file_path': file_path,
                'file_type': file_type,
                'status': 'success',
                'content_length': len(content_text) if content_text else 0,
                'metadata': metadata,
                'course_metadata': course_metadata
            }

        except Exception as e:
            self.logger.error(f"解析文件 {file_path} 时出错: {str(e)}")
            return {
                'content_text': '',
                'file_path': file_path,
                'file_type': file_type,
                'status': 'error',
                'error_message': str(e)
            }

    def _get_empty_course_metadata(self) -> Dict[str, str]:
        """
        返回空的课程元数据字典
        """
        return {
            'course_topic': '',
            'core_knowledge': '',
            'learning_objectives': '',
            'application_scenarios': '',
            'difficulty_level': ''
        }

    def _generate_course_metadata(self, content_text: str) -> Dict[str, Any]:
        """
        使用DeepSeek API生成课程元数据

        Args:
            content_text: 课程内容文本

        Returns:
            课程元数据字典
        """
        if not self.deepseek_client:
            raise ValueError("DeepSeek API客户端未初始化")

        try:
            # 限制内容长度以避免token限制
            if len(content_text) > 4000:
                content_preview = content_text[:4000] + "...[内容已截断]"
            else:
                content_preview = content_text

            response = self.deepseek_client.chat.completions.create(
                model="deepseek-chat",
                messages=[
                    {
                        "role": "system",
                        "content": """
                        你是一位教育专家，请对提供的大数据课程课件进行专业概述。
                        请严格按照以下结构输出，只输出这五个部分，不要额外内容：
                        1. 课程主题概括（2-3句话）
                        2. 核心知识点（列出3-5个关键概念）
                        3. 学习目标（学生将掌握什么能力）
                        4. 实际应用场景
                        5. 难度级别评估
                        要求：语言简洁专业，重点突出，适合学生快速了解课程内容。
                        """
                    },
                    {
                        "role": "user",
                        "content": f"课件内容：{content_preview}"
                    },
                ],
                stream=False
            )

            result_text = response.choices[0].message.content
            self.logger.info("DeepSeek API响应成功")

            # 解析结果文本，提取五个部分
            return self._parse_course_metadata(result_text)

        except Exception as e:
            self.logger.error(f"DeepSeek API调用失败: {str(e)}")
            # 返回空的课程元数据而不是抛出异常
            return self._get_empty_course_metadata()

    def _parse_course_metadata(self, result_text: str) -> Dict[str, str]:
        """
        解析DeepSeek API返回的课程元数据

        Args:
            result_text: API返回的文本

        Returns:
            解析后的课程元数据字典
        """
        course_metadata = self._get_empty_course_metadata()

        try:
            lines = result_text.strip().split('\n')
            current_section = None
            section_content = []

            for line in lines:
                line = line.strip()
                if not line:
                    continue

                # 检查是否是新的章节标题
                if line.startswith('1.') or '课程主题概括' in line:
                    if current_section and section_content:
                        course_metadata[current_section] = '\n'.join(section_content).strip()
                    current_section = 'course_topic'
                    section_content = [self._clean_section_line(line, '1.', '课程主题概括')]
                elif line.startswith('2.') or '核心知识点' in line:
                    if current_section and section_content:
                        course_metadata[current_section] = '\n'.join(section_content).strip()
                    current_section = 'core_knowledge'
                    section_content = [self._clean_section_line(line, '2.', '核心知识点')]
                elif line.startswith('3.') or '学习目标' in line:
                    if current_section and section_content:
                        course_metadata[current_section] = '\n'.join(section_content).strip()
                    current_section = 'learning_objectives'
                    section_content = [self._clean_section_line(line, '3.', '学习目标')]
                elif line.startswith('4.') or '实际应用场景' in line:
                    if current_section and section_content:
                        course_metadata[current_section] = '\n'.join(section_content).strip()
                    current_section = 'application_scenarios'
                    section_content = [self._clean_section_line(line, '4.', '实际应用场景')]
                elif line.startswith('5.') or '难度级别评估' in line:
                    if current_section and section_content:
                        course_metadata[current_section] = '\n'.join(section_content).strip()
                    current_section = 'difficulty_level'
                    section_content = [self._clean_section_line(line, '5.', '难度级别评估')]
                elif current_section:
                    section_content.append(line)

            # 处理最后一个章节
            if current_section and section_content:
                course_metadata[current_section] = '\n'.join(section_content).strip()

            return course_metadata

        except Exception as e:
            self.logger.error(f"解析课程元数据失败: {str(e)}")
            # 返回原始文本作为备用
            return {'raw_course_metadata': result_text}

    def _clean_section_line(self, line: str, prefix: str, section_name: str) -> str:
        """
        清理章节行，移除前缀和章节名称
        """
        line = line.replace(prefix, '').replace(section_name, '').strip()
        # 移除可能的分隔符
        line = line.strip(':-').strip()
        return line

    def _extract_specific_metadata(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """
        提取特定文件类型的元数据

        Args:
            file_path: 文件路径
            file_type: 文件类型

        Returns:
            特定文件类型的元数据
        """
        metadata = {}

        try:
            if file_type in ['pdf']:
                # PDF文件元数据
                with pdfplumber.open(file_path) as pdf:
                    metadata['pages'] = len(pdf.pages)

            elif file_type in ['ppt', 'pptx']:
                # PPT文件元数据
                presentation = Presentation(file_path)
                metadata['slides'] = len(presentation.slides)

            elif file_type in ['docx']:
                # DOCX文件元数据
                doc = Document(file_path)
                metadata['paragraphs'] = len(doc.paragraphs)
                metadata['tables'] = len(doc.tables)

            elif file_type in ['jpg', 'jpeg', 'png']:
                # 图片文件元数据
                with Image.open(file_path) as img:
                    metadata['image_width'] = img.width
                    metadata['image_height'] = img.height
                    metadata['image_mode'] = img.mode

            elif file_type in ['video', 'mp4', 'avi', 'mov', 'mkv']:
                # 视频文件元数据
                duration = self._get_video_duration(file_path)
                metadata['duration'] = duration
                metadata['duration_formatted'] = self._format_duration(duration)

        except Exception as e:
            self.logger.warning(f"提取特定元数据失败: {str(e)}")

        return metadata

    def _get_video_duration(self, video_path: str) -> float:
        """
        获取视频时长

        Args:
            video_path: 视频文件路径

        Returns:
            视频时长（秒）
        """
        try:
            # 使用OpenCV获取视频时长
            cap = cv2.VideoCapture(video_path)
            if cap.isOpened():
                fps = cap.get(cv2.CAP_PROP_FPS)
                frame_count = cap.get(cv2.CAP_PROP_FRAME_COUNT)
                duration = frame_count / fps if fps > 0 else 0
                cap.release()
                return duration
            return 0
        except Exception as e:
            self.logger.warning(f"获取视频时长失败: {str(e)}")
            return 0

    def _format_duration(self, seconds: float) -> str:
        """
        格式化时长

        Args:
            seconds: 秒数

        Returns:
            格式化的时长字符串
        """
        if seconds == 0:
            return "未知"

        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)

        if hours > 0:
            return f"{hours:02d}:{minutes:02d}:{secs:02d}"
        else:
            return f"{minutes:02d}:{secs:02d}"

    def _parse_image(self, file_path: str) -> str:
        """解析图片文件"""
        try:
            pipeline = PPStructureV3(
                use_doc_orientation_classify=False,
                use_doc_unwarping=False
            )

            # For Image
            output = pipeline.predict(
                input=file_path,
            )

            # Visualize the results and save the JSON results
            for res in output:
                res.save_to_markdown(save_path=".")

            filename = os.path.basename(file_path)
            # 去除扩展名
            filename_without_extension = os.path.splitext(filename)[0]
            file_md_path = filename_without_extension + ".md"
            with open(file_md_path, 'r', encoding='utf-8') as file:
                content = file.read()
            os.remove(file_md_path)
            return content

        except Exception as e:
            self.logger.error(f"图片解析错误: {str(e)}")
            raise

    def _parse_pdf(self, file_path: str) -> str:
        """解析PDF文件"""
        try:
            result = extract_file_sync(file_path)
            return result.content
        except Exception as e:
            self.logger.error(f"PDF解析错误: {str(e)}")
            raise

    def _parse_ppt(self, file_path: str) -> str:
        """解析PPT/PPTX文件"""
        try:
            result = extract_file_sync(file_path)
            return result.content
        except Exception as e:
            self.logger.error(f"PPT解析错误: {str(e)}")
            raise

    def _parse_docx(self, file_path: str) -> str:
        """解析DOCX文件"""
        try:
            result = extract_file_sync(file_path)
            return result.content
        except Exception as e:
            self.logger.error(f"DOCX解析错误: {str(e)}")
            raise

    def _parse_video(self, file_path: str) -> str:
        """解析视频文件，提取音频并转换为文本"""
        try:
            # 提取音频
            audio_path = self._extract_audio_from_video(file_path)

            # 语音识别
            text_content = self._speech_to_text(audio_path)

            # 清理临时文件
            if os.path.exists(audio_path):
                os.remove(audio_path)

            return text_content

        except Exception as e:
            self.logger.error(f"视频解析错误: {str(e)}")
            raise

    def _extract_audio_from_video(self, video_path: str) -> str:
        """从视频中提取音频"""
        try:
            temp_audio_path = 'example.mp3'
            subprocess.run(
                ['ffmpeg', '-i', video_path, '-c:a', 'libmp3lame', '-b:a',
                 '192k',
                 temp_audio_path])

            return temp_audio_path

        except Exception as e:
            self.logger.error(f"音频提取错误: {str(e)}")
            raise

    def _speech_to_text(self, audio_path: str) -> str:
        """将音频转换为文本"""
        try:
            result = model.transcribe(audio_path)
            print(result["text"])

            return ''.join(result["text"])

        except Exception as e:
            self.logger.error(f"语音识别错误: {str(e)}")
            raise

    def update_content_field(self, data: Dict[str, Any], generate_course_metadata: bool = False) -> Dict[str, Any]:
        """
        更新数据的content_text字段

        Args:
            data: 包含文件信息的数据字典
            generate_course_metadata: 是否生成课程元数据

        Returns:
            更新后的数据字典
        """
        if 'file_path' not in data or 'file_type' not in data:
            raise ValueError("数据中必须包含file_path和file_type字段")

        result = self.extract_content(data['file_path'], data['file_type'], generate_course_metadata)

        # 保存提取的文本到文件（用于调试）
        with open("extracted_text.txt", "w", encoding="utf-8") as output_file:
            output_file.write(result['content_text'])

        # 更新content_text字段和元数据
        data['content_text'] = result['content_text']
        data['parse_status'] = result['status']
        data['metadata'] = result.get('metadata', {})

        # 添加课程元数据
        if generate_course_metadata:
            data['course_metadata'] = result.get('course_metadata', {})

        if result['status'] == 'error':
            data['error'] = result.get('error_message', '未知错误')

        return data

    def get_file_metadata_only(self, file_path: str) -> Dict[str, Any]:
        """
        仅获取文件元数据，不提取内容

        Args:
            file_path: 文件路径

        Returns:
            文件元数据字典
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"文件不存在: {file_path}")

            # 提取基本元数据
            metadata = self.extract_metadata(file_path)

            # 提取文件类型
            file_extension = Path(file_path).suffix.lower()
            file_type = file_extension[1:] if file_extension.startswith('.') else file_extension

            # 根据文件类型提取特定元数据
            specific_metadata = self._extract_specific_metadata(file_path, file_type)
            metadata.update(specific_metadata)

            return {
                'status': 'success',
                'file_path': file_path,
                'file_type': file_type,
                'metadata': metadata
            }

        except Exception as e:
            self.logger.error(f"获取文件元数据失败: {str(e)}")
            return {
                'status': 'error',
                'file_path': file_path,
                'error_message': str(e)
            }


# 使用示例
if __name__ == "__main__":
    # 初始化解析器（需要提供DeepSeek API密钥）
    parser = FileParser(deepseek_api_key="sk-0e7d09c913e6426dbf7e55450b67daa5")

    # 示例：获取文件元数据
    sample_file = '3-1-cn.docx'  # 替换为实际文件路径

    try:
        # 仅获取元数据
        metadata_result = parser.get_file_metadata_only(sample_file)
        if metadata_result['status'] == 'success':
            print("文件元数据:")
            for key, value in metadata_result['metadata'].items():
                print(f"  {key}: {value}")

        print("\n" + "=" * 50 + "\n")

        # 获取内容和元数据（包括课程元数据）
        sample_data = {
            'file_path': sample_file,
            'file_type': metadata_result['metadata']['file_type'],
            'content_text': ''  # 将被更新
        }

        # 提取内容并更新字段，同时生成课程元数据
        updated_data = parser.update_content_field(sample_data, generate_course_metadata=True)
        print(f"提取成功，内容长度: {len(updated_data['content_text'])}")
        print(f"状态: {updated_data['parse_status']}")

        # 输出课程元数据
        if 'course_metadata' in updated_data:
            print("\n课程元数据:")
            course_meta = updated_data['course_metadata']
            print(f"  课程主题: {course_meta.get('course_topic', '无')}")
            print(f"  核心知识点: {course_meta.get('core_knowledge', '无')}")
            print(f"  学习目标: {course_meta.get('learning_objectives', '无')}")
            print(f"  实际应用场景: {course_meta.get('application_scenarios', '无')}")
            print(f"  难度级别: {course_meta.get('difficulty_level', '无')}")

    except Exception as e:
        print(f"处理失败: {e}")